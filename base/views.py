#  Authors: 
#  Franyelyn Rodríguez
#  Christopher Gedler <cgedler@gmail.com>
#  File name: views.py
#  Version: 1.0
#  Last Change: 29.06.2022 21:07:30 -04
#  Description: Vistas del Sistema Midipaaa

# Imports base para el sistema:
from django.shortcuts import get_object_or_404, render
from django.http import Http404, HttpResponse, HttpResponseRedirect
from django.urls import reverse
from django.views import generic
from django.views.generic import View
from django.views.generic.edit import CreateView, DeleteView, UpdateView
from django.contrib import messages
from django.contrib.messages.views import SuccessMessageMixin
from django.db.models import Q

# Imports Modelos:
from .models import *
from .forms import *
# Export PDF:
from .render_to_pdf import *
# Export Word:
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
import os
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
# Export Excel:
import openpyxl
from openpyxl import Workbook
from openpyxl.workbook import Workbook
from openpyxl.cell import WriteOnlyCell
from openpyxl.comments import Comment
from openpyxl.styles import Font, Color, Fill, Alignment, PatternFill
from openpyxl.styles.borders import Border, Side

EMPRESA_NOMBRE = 'Fundación Familia Midipaaa'
EMPRESA_RIF = 'J-0000000-0'
EMPRESA_DIRESTADO = 'Miranda'
EMPRESA_DIRCIUDAD = 'Los Teques'
EMPRESA_DIRECCION = 'Santa Eulalia'
EMPRESA_POSTAL = '1201'


def index(request):
    """
    Description: Función que devuelve la vista inicial para entrar al
    Sistema Midipaaa
    Return:
        Pagina Index principal
    """
    return render(request, 'base/index.html')


def DevolverCount(query):
    total = query.count()
    return total


def DevolverData_BeneficiadoSector(listado):
    data = []
    if len(listado) < 1:
        data = [0,0,0]
        return data
    else:
        for item in listado:
            data_encontrada = DevolverCount(
            Beneficiado.objects.filter(Q(sector=str(item))))
            data.append(data_encontrada)
        return data


def DevolverData_BeneficiadoCurso(listado):
    data = []
    if len(listado) < 1:
        data = [0,0,0]
        return data
    else:
        for item in listado:
            data_encontrada = DevolverCount(
            CursoBeneficiado.objects.filter(Q(curso=str(item))))
            data.append(data_encontrada)
        return data


def DevolverData_BeneficiadoCaja(listado):
    data = []
    if len(listado) < 1:
        data = [0,0,0]
        return data
    else:
        for item in listado:
            data_encontrada = DevolverCount(
            CajaBeneficiado.objects.filter(Q(caja=str(item))))
            data.append(data_encontrada)
        return data


def home(request):
    """
    Description: Función que devuelve la vista home con los gráficos
    Sistema Midipaaa
    Return:
        Pagina Index principal
    """
    labels_sectores = ['1', '2', '3', '4', '5', '6', '7', '8', '9']
    data_sectores = DevolverData_BeneficiadoSector(labels_sectores)

    labels_cursos = ['1', '2', '3', '4', '5', '6']
    data_cursos = DevolverData_BeneficiadoCurso(labels_cursos)

    labels_cajas = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10']
    data_cajas = DevolverData_BeneficiadoCaja(labels_cajas)

    return render(request, 'base/home.html', {
        'labels_sector': labels_sectores,
        'data_sector': data_sectores,
        'labels_curso': labels_cursos,
        'data_curso': data_cursos,
        'labels_caja': labels_cajas,
        'data_caja': data_cajas,
        })


def about(request):
    return HttpResponse('Acerca de...')


def contact(request):
    return HttpResponse('Datos de contacto')


# Beneficiado:
class BeneficiadoList(generic.ListView):
    model = Beneficiado
    template_name = 'base/Beneficiado_i.html'
    queryset = Beneficiado.objects.all().order_by('-created')
    context_object_name = 'Beneficiado'
    extra_context = {
        'section_title': 'Listado de Beneficiados',
        'module_title': 'Beneficiado'
    }


class BeneficiadoCreate(SuccessMessageMixin, CreateView):
    model = Beneficiado
    form_class = BeneficiadoForm
    extra_context = {
        'section_title': 'Crear Beneficiado',
        'module_title': 'Beneficiado',
    }
    template_name = 'base/Beneficiado_c.html'
    success_message = 'Creado Correctamente!'

    def get_success_url(self):
        return reverse('base:BeneficiadoList')


class BeneficiadoDetail(generic.DetailView):
    model = Beneficiado
    form_class = BeneficiadoForm
    template_name = 'base/Beneficiado_d.html'

    def get_context_data(self, **kwargs):
        context = super(BeneficiadoDetail, self).get_context_data(**kwargs)
        context['now'] = timezone.now()
        extra_context = {
            'section_title': 'Detalle del Beneficiado',
            'module_title': 'Beneficiado'
        }
        return context


class BeneficiadoUpdate(SuccessMessageMixin, UpdateView):
    model = Beneficiado
    form_class = BeneficiadoForm
    template_name = 'base/Beneficiado_u.html'
    success_message = 'Actualizado Correctamente!'

    def get_success_url(self):
        return reverse('base:BeneficiadoList')


class BeneficiadoDelete(SuccessMessageMixin, DeleteView):
    model = Beneficiado
    form = Beneficiado
    fields = '__all__'

    def get_success_url(self):
        success_message = 'Eliminado Correctamente!'
        messages.success(self.request, (success_message))
        return reverse('base:BeneficiadoList')


class BeneficiadoPdfList(View):
    def get(self, request, *args, **kwargs):
        now = timezone.now()
        data = {
            'section_title': 'Listado de Beneficiados',
            'module_title': 'Beneficiado',
            'today': now,
            'empresa_nombre': EMPRESA_NOMBRE,
            'empresa_rif': EMPRESA_RIF,
            'empresa_direstado': EMPRESA_DIRESTADO,
            'empresa_dirciudad': EMPRESA_DIRCIUDAD,
            'empresa_direccion': EMPRESA_DIRECCION,
            'empresa_postal': EMPRESA_POSTAL,
            'context_pdf': Beneficiado.objects.all()
        }
        pdf = render_to_pdf('base/BeneficiadoPdfList.html', data)
        return HttpResponse(pdf, content_type='application/pdf')


class BeneficiadoPdfUnid(View):
    def get(self, request, *args, **kwargs):
        id = str
        id = kwargs['pk']
        object_list = Beneficiado.objects.filter(id_beneficiado__icontains=id)
        now = timezone.now()
        data = {
            'section_title': 'Datos del Beneficiado',
            'module_title': 'Datos del Beneficiado',
            'today': now,
            'empresa_nombre': EMPRESA_NOMBRE,
            'empresa_rif': EMPRESA_RIF,
            'empresa_direstado': EMPRESA_DIRESTADO,
            'empresa_dirciudad': EMPRESA_DIRCIUDAD,
            'empresa_direccion': EMPRESA_DIRECCION,
            'empresa_postal': EMPRESA_POSTAL,
            'context_pdf': object_list
        }
        pdf = render_to_pdf('base/BeneficiadoPdfUnid.html', data)
        return HttpResponse(pdf, content_type='application/pdf')


class BeneficiadoWord(View):
    def get(self, request, *args, **kwargs):
        id = str
        id = kwargs['pk']
        results = Beneficiado.objects.filter(id_beneficiado__exact=id)
        date = datetime.now().strftime('%d-%m-%Y')
        document = Document(os.path.join(BASE_DIR, 'base/base.docx'))
        style_normal = document.styles['Normal']
        font = style_normal.font
        font.name = 'Arial'
        font.size = Pt(11)
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(1.0)
            section.bottom_margin = Cm(1.0)
            section.left_margin = Cm(1.0)
            section.right_margin = Cm(1.0)
        table_head = document.add_table(rows=3, cols=2)
        row_1 = table_head.rows[0]
        row_1.cells[0].text = EMPRESA_NOMBRE
        formato_emp = row_1.cells[0].paragraphs
        for paragraph in formato_emp:
            for run in paragraph.runs:
                font = run.font
                font.name = 'Arial'
                font.size = Pt(16)
        row_1.cells[1].text = '\t\t\t\tFecha : ' + date
        row_2 = table_head.rows[1]
        row_2.cells[0].text = EMPRESA_RIF
        row_3 = table_head.rows[2]
        row_3.cells[0].text = EMPRESA_DIRCIUDAD + ' ' + EMPRESA_DIRESTADO + ' ' + EMPRESA_POSTAL
        document.add_paragraph()
        title = document.add_paragraph('Datos del Beneficiado')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph('Cédula : ' + str(results[0].id_beneficiado))
        document.add_paragraph('Nombres : ' +  results[0].primer_nombre + ',' + results[0].segundo_nombre)
        document.add_paragraph('Apellidos : ' + results[0].primer_apellido + ',' + results[0].segundo_apellido)
        sexo =  results[0].get_sexo_display()
        document.add_paragraph('Sexo : ' + str(sexo))
        nacionalidad =  results[0].get_nacionalidad_display()
        document.add_paragraph('Nacionalidad : ' + str(nacionalidad))
        estado_civil =  results[0].get_estado_civil_display()
        document.add_paragraph('Estado civil : ' + str(estado_civil))
        document.add_paragraph('Fecha de nacimiento : ' + str(results[0].fecha_nacimiento))
        document.add_paragraph('Dirección : ' + results[0].direccion)
        document.add_paragraph('Celular : ' + results[0].celular)
        document.add_paragraph('Teléfono : ' + results[0].telefono)
        document.add_paragraph('Email : ' + results[0].email)
        document.add_paragraph('Sector : ' + str(results[0].sector))
        document.add_paragraph('Cantidad de familiares : ' + str(results[0].familiares))
        ingreso =  results[0].get_ingreso_display()
        document.add_paragraph('Ingreso familiar : ' + str(ingreso))
        estudio =  results[0].get_estudio_display()
        document.add_paragraph('Nivel de estudio : ' + str(estudio))
        vivienda =  results[0].get_vivienda_display()
        document.add_paragraph('Tipo de vivienda : ' + str(vivienda))
        tenencia =  results[0].get_tenencia_display()
        document.add_paragraph('Tenencia de la vivienda : ' + str(tenencia)) 
        document.add_page_break()
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename={date}-Beneficiado.docx'.format(
            date=datetime.now().strftime('%d-%m-%Y'),
        )
        document.save(response)
        return response


class BeneficiadoExcel(View):
    def get(self, request, *args, **kwargs):
        id = str
        id = kwargs['pk']
        results = Beneficiado.objects.filter(id_beneficiado__exact=id)
        date = datetime.now().strftime('%d-%m-%Y')
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
        response['Content-Disposition'] = 'attachment; filename={date}-Beneficiado.xlsx'.format(
            date=datetime.now().strftime('%d-%m-%Y'),
        )
        sexo =  results[0].get_sexo_display()
        nacionalidad =  results[0].get_nacionalidad_display()
        estado_civil =  results[0].get_estado_civil_display()
        ingreso =  results[0].get_ingreso_display()
        estudio =  results[0].get_estudio_display()
        vivienda =  results[0].get_vivienda_display()
        tenencia =  results[0].get_tenencia_display()
        wb = Workbook()
        ws = wb.active
        ws.column_dimensions['A'].width = 2.0
        ws.column_dimensions['B'].width = 25.0
        ws.column_dimensions['C'].width = 25.0
        ws.column_dimensions['D'].width = 25.0
        ws['B1'] = EMPRESA_NOMBRE
        ws['B2'] = EMPRESA_RIF
        ws['B3'] = EMPRESA_DIRCIUDAD + ' ' + EMPRESA_DIRESTADO
        ws['B4'] = EMPRESA_POSTAL
        ws['D1'] = date
        ws['C5'] = 'Datos del Beneficiado'
        ws['B7'] = 'Cédula :'
        ws['B8'] = 'Nombres :'
        ws['B9'] = 'Apellidos :'
        ws['B10'] = 'Sexo :'
        ws['B11'] = 'Nacionalidad :'
        ws['B12'] = 'Estado civil :'
        ws['B13'] = 'Fecha de nacimiento :'
        ws['B14'] = 'Dirección :'
        ws['B15'] = 'Celular :'
        ws['B16'] = 'Teléfono :'
        ws['B17'] = 'Email :'
        ws['B18'] = 'Sector :'
        ws['B19'] = 'Cantidad de familiares :'
        ws['B20'] = 'Ingreso familiar :'
        ws['B21'] = 'Nivel de estudio :'
        ws['B22'] = 'Tipo de vivienda :'
        ws['B23'] = 'Tenencia de la vivienda :'
        ws['C7'] = str(results[0].id_beneficiado)
        ws['C8'] = results[0].primer_nombre + ',' + results[0].segundo_nombre
        ws['C9'] = results[0].primer_apellido + ',' + results[0].segundo_apellido
        ws['C10'] = str(sexo)
        ws['C11'] = str(nacionalidad)
        ws['C12'] = str(estado_civil)
        ws['C13'] = str(results[0].fecha_nacimiento)
        ws['C14'] = results[0].direccion
        ws['C15'] = results[0].celular
        ws['C16'] = results[0].telefono
        ws['C17'] = results[0].email
        ws['C18'] = str(results[0].sector)
        ws['C19'] = str(results[0].familiares)
        ws['C20'] = str(ingreso)
        ws['C21'] = str(estudio)
        ws['C22'] = str(vivienda)
        ws['C23'] = str(tenencia)
        wb.save(response)
        return response


# Asignar Caja Beneficiado:
class AsignarCajaList(generic.ListView):
    model = CajaBeneficiado
    template_name = 'base/AsignarCaja_i.html'
    queryset = CajaBeneficiado.objects.all().order_by('-created')
    context_object_name = 'AsignarCaja'
    extra_context = {
        'section_title': 'Listado de Cajas Asignadas',
        'module_title': 'Asignar Caja'
    }


class AsignarCajaCreate(SuccessMessageMixin, CreateView):
    model = CajaBeneficiado
    form_class = AsignarCajaForm
    extra_context = {
        'section_title': 'Crear Asignar Caja',
        'module_title': 'Asignar Caja',
    }
    template_name = 'base/AsignarCaja_c.html'
    success_message = 'Creado Correctamente!'

    def get_success_url(self):
        return reverse('base:AsignarCajaList')


class AsignarCajaDetail(generic.DetailView):
    model = CajaBeneficiado
    form_class = AsignarCajaForm
    template_name = 'base/AsignarCaja_d.html'

    def get_context_data(self, **kwargs):
        context = super(AsignarCajaDetail, self).get_context_data(**kwargs)
        context['now'] = timezone.now()
        extra_context = {
            'section_title': 'Detalle de Caja Asignada',
            'module_title': 'Asignar Caja'
        }
        return context


class AsignarCajaUpdate(SuccessMessageMixin, UpdateView):
    model = CajaBeneficiado
    form_class = AsignarCajaForm
    template_name = 'base/AsignarCaja_u.html'
    success_message = 'Actualizado Correctamente!'

    def get_success_url(self):
        return reverse('base:AsignarCajaList')


class AsignarCajaDelete(SuccessMessageMixin, DeleteView):
    model = CajaBeneficiado
    form = AsignarCajaForm
    fields = '__all__'

    def get_success_url(self):
        success_message = 'Eliminado Correctamente!'
        messages.success(self.request, (success_message))
        return reverse('base:AsignarCajaList')


class AsignarCajaPdfList(View):
    def get(self, request, *args, **kwargs):
        now = timezone.now()
        data = {
            'section_title': 'Listado de Cajas Asignadas',
            'module_title': 'Asignar Caja',
            'today': now,
            'empresa_nombre': EMPRESA_NOMBRE,
            'empresa_rif': EMPRESA_RIF,
            'empresa_direstado': EMPRESA_DIRESTADO,
            'empresa_dirciudad': EMPRESA_DIRCIUDAD,
            'empresa_direccion': EMPRESA_DIRECCION,
            'empresa_postal': EMPRESA_POSTAL,
            'context_pdf': CajaBeneficiado.objects.all()
        }
        pdf = render_to_pdf('base/AsignarCajaPdfList.html', data)
        return HttpResponse(pdf, content_type='application/pdf')


# Asignar Curso Beneficiado:
class AsignarCursoList(generic.ListView):
    model = CursoBeneficiado
    template_name = 'base/AsignarCurso_i.html'
    queryset = CursoBeneficiado.objects.all().order_by('-created')
    context_object_name = 'AsignarCurso'
    extra_context = {
        'section_title': 'Listado de Cursos Asignados',
        'module_title': 'Asignar Curso'
    }


class AsignarCursoCreate(SuccessMessageMixin, CreateView):
    model = CursoBeneficiado
    form_class = AsignarCursoForm
    extra_context = {
        'section_title': 'Crear Asignar Curso',
        'module_title': 'Asignar Curso',
    }
    template_name = 'base/AsignarCurso_c.html'
    success_message = 'Creado Correctamente!'

    def get_success_url(self):
        return reverse('base:AsignarCursoList')


class AsignarCursoDetail(generic.DetailView):
    model = CursoBeneficiado
    form_class = AsignarCursoForm
    template_name = 'base/AsignarCurso_d.html'

    def get_context_data(self, **kwargs):
        context = super(AsignarCursoDetail, self).get_context_data(**kwargs)
        context['now'] = timezone.now()
        extra_context = {
            'section_title': 'Detalle de Curso Asignado',
            'module_title': 'Asignar Curso'
        }
        return context


class AsignarCursoUpdate(SuccessMessageMixin, UpdateView):
    model = CursoBeneficiado
    form_class = AsignarCursoForm
    template_name = 'base/AsignarCurso_u.html'
    success_message = 'Actualizado Correctamente!'

    def get_success_url(self):
        return reverse('base:AsignarCursoList')


class AsignarCursoDelete(SuccessMessageMixin, DeleteView):
    model = CursoBeneficiado
    form = AsignarCursoForm
    fields = '__all__'

    def get_success_url(self):
        success_message = 'Eliminado Correctamente!'
        messages.success(self.request, (success_message))
        return reverse('base:AsignarCursoList')


class AsignarCursoPdfList(View):
    def get(self, request, *args, **kwargs):
        now = timezone.now()
        data = {
            'section_title': 'Listado de Cursos Asignados',
            'module_title': 'Asignar Curso',
            'today': now,
            'empresa_nombre': EMPRESA_NOMBRE,
            'empresa_rif': EMPRESA_RIF,
            'empresa_direstado': EMPRESA_DIRESTADO,
            'empresa_dirciudad': EMPRESA_DIRCIUDAD,
            'empresa_direccion': EMPRESA_DIRECCION,
            'empresa_postal': EMPRESA_POSTAL,
            'context_pdf': CursoBeneficiado.objects.all()
        }
        pdf = render_to_pdf('base/AsignarCursoPdfList.html', data)
        return HttpResponse(pdf, content_type='application/pdf')


# Caja:
class CajaList(generic.ListView):
    model = Caja
    template_name = 'base/Caja_i.html'
    queryset = Caja.objects.all().order_by('-created')
    context_object_name = 'Caja'
    extra_context = {
        'section_title': 'Listado de Cajas',
        'module_title': 'Caja'
    }


class CajaCreate(SuccessMessageMixin, CreateView):
    model = Caja
    form_class = CajaForm
    extra_context = {
        'section_title': 'Crear Caja',
        'module_title': 'Caja',
    }
    template_name = 'base/Caja_c.html'
    success_message = 'Creado Correctamente!'

    def get_success_url(self):
        return reverse('base:CajaList')


class CajaDetail(generic.DetailView):
    model = Caja
    form_class = CajaForm
    template_name = 'base/Caja_d.html'

    def get_context_data(self, **kwargs):
        context = super(CajaDetail, self).get_context_data(**kwargs)
        context['now'] = timezone.now()
        extra_context = {
            'section_title': 'Detalle de la Caja',
            'module_title': 'Caja'
        }
        return context


class CajaUpdate(SuccessMessageMixin, UpdateView):
    model = Caja
    form_class = CajaForm
    template_name = 'base/Caja_u.html'
    success_message = 'Actualizado Correctamente!'

    def get_success_url(self):
        return reverse('base:CajaList')


class CajaDelete(SuccessMessageMixin, DeleteView):
    model = Caja
    form = Caja
    fields = '__all__'

    def get_success_url(self):
        success_message = 'Eliminado Correctamente!'
        messages.success(self.request, (success_message))
        return reverse('base:CajaList')


class CajaPdfList(View):
    def get(self, request, *args, **kwargs):
        now = timezone.now()
        data = {
            'section_title': 'Listado de Cajas',
            'module_title': 'Listado de Cajas',
            'today': now,
            'empresa_nombre': EMPRESA_NOMBRE,
            'empresa_rif': EMPRESA_RIF,
            'empresa_direstado': EMPRESA_DIRESTADO,
            'empresa_dirciudad': EMPRESA_DIRCIUDAD,
            'empresa_direccion': EMPRESA_DIRECCION,
            'empresa_postal': EMPRESA_POSTAL,
            'context_pdf': Caja.objects.all()
        }
        pdf = render_to_pdf('base/CajaPdfList.html', data)
        return HttpResponse(pdf, content_type='application/pdf')


class CajaPdfUnid(View):
    def get(self, request, *args, **kwargs):
        id = str
        id = kwargs['pk']
        object_list = Caja.objects.filter(id__icontains=id)
        now = timezone.now()
        data = {
            'section_title': 'Datos de la Caja',
            'module_title': 'Datos de la Caja',
            'today': now,
            'empresa_nombre': EMPRESA_NOMBRE,
            'empresa_rif': EMPRESA_RIF,
            'empresa_direstado': EMPRESA_DIRESTADO,
            'empresa_dirciudad': EMPRESA_DIRCIUDAD,
            'empresa_direccion': EMPRESA_DIRECCION,
            'empresa_postal': EMPRESA_POSTAL,
            'context_pdf': object_list
        }
        pdf = render_to_pdf('base/CajaPdfUnid.html', data)
        return HttpResponse(pdf, content_type='application/pdf')


class CajaWord(View):
    def get(self, request, *args, **kwargs):
        id = str
        id = kwargs['pk']
        results = Caja.objects.filter(id__exact=id)
        date = datetime.now().strftime('%d-%m-%Y')
        document = Document(os.path.join(BASE_DIR, 'base/base.docx'))
        style_normal = document.styles['Normal']
        font = style_normal.font
        font.name = 'Arial'
        font.size = Pt(11)
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(1.0)
            section.bottom_margin = Cm(1.0)
            section.left_margin = Cm(1.0)
            section.right_margin = Cm(1.0)
        table_head = document.add_table(rows=3, cols=2)
        row_1 = table_head.rows[0]
        row_1.cells[0].text = EMPRESA_NOMBRE
        formato_emp = row_1.cells[0].paragraphs
        for paragraph in formato_emp:
            for run in paragraph.runs:
                font = run.font
                font.name = 'Arial'
                font.size = Pt(16)
        row_1.cells[1].text = '\t\t\t\tFecha : ' + date
        row_2 = table_head.rows[1]
        row_2.cells[0].text = EMPRESA_RIF
        row_3 = table_head.rows[2]
        row_3.cells[0].text = EMPRESA_DIRCIUDAD + ' ' + EMPRESA_DIRESTADO + ' ' + EMPRESA_POSTAL
        document.add_paragraph()
        title = document.add_paragraph('Datos de la Caja')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph('ID : ' + str(results[0].id))
        document.add_paragraph('Descripción : ' +  results[0].descripcion)
        document.add_paragraph('Artículos : ' + results[0].articulos)
        document.add_page_break()
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename={date}-Caja.docx'.format(
            date=datetime.now().strftime('%d-%m-%Y'),
        )
        document.save(response)
        return response


class CajaExcel(View):
    def get(self, request, *args, **kwargs):
        id = str
        id = kwargs['pk']
        results = Caja.objects.filter(id__exact=id)
        date = datetime.now().strftime('%d-%m-%Y')
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
        response['Content-Disposition'] = 'attachment; filename={date}-Caja.xlsx'.format(
            date=datetime.now().strftime('%d-%m-%Y'),
        )
        wb = Workbook()
        ws = wb.active
        ws.column_dimensions['A'].width = 2.0
        ws.column_dimensions['B'].width = 25.0
        ws.column_dimensions['C'].width = 25.0
        ws.column_dimensions['D'].width = 25.0
        ws['B1'] = EMPRESA_NOMBRE
        ws['B2'] = EMPRESA_RIF
        ws['B3'] = EMPRESA_DIRCIUDAD + ' ' + EMPRESA_DIRESTADO
        ws['B4'] = EMPRESA_POSTAL
        ws['D1'] = date
        ws['C5'] = 'Datos de la Caja'
        ws['B7'] = 'ID :'
        ws['B8'] = 'Descripción :'
        ws['B9'] = 'Artículos :'
        ws['C7'] = str(results[0].id)
        ws['C8'] = results[0].descripcion
        ws['C9'] = results[0].articulos
        wb.save(response)
        return response


# Curso:
class CursoList(generic.ListView):
    model = Curso
    template_name = 'base/Curso_i.html'
    queryset = Curso.objects.all().order_by('-created')
    context_object_name = 'Curso'
    extra_context = {
        'section_title': 'Listado de Cursos',
        'module_title': 'Curso'
    }


class CursoCreate(SuccessMessageMixin, CreateView):
    model = Curso
    form_class = CursoForm
    extra_context = {
        'section_title': 'Crear Curso',
        'module_title': 'Curso',
    }
    template_name = 'base/Curso_c.html'
    success_message = 'Creado Correctamente!'

    def get_success_url(self):
        return reverse('base:CursoList')


class CursoDetail(generic.DetailView):
    model = Curso
    form_class = CursoForm
    template_name = 'base/Curso_d.html'

    def get_context_data(self, **kwargs):
        context = super(CursoDetail, self).get_context_data(**kwargs)
        context['now'] = timezone.now()
        extra_context = {
            'section_title': 'Detalle del Curso',
            'module_title': 'Curso'
        }
        return context


class CursoUpdate(SuccessMessageMixin, UpdateView):
    model = Curso
    form_class = CursoForm
    template_name = 'base/Curso_u.html'
    success_message = 'Actualizado Correctamente!'

    def get_success_url(self):
        return reverse('base:CursoList')


class CursoDelete(SuccessMessageMixin, DeleteView):
    model = Curso
    form = Curso
    fields = '__all__'

    def get_success_url(self):
        success_message = 'Eliminado Correctamente!'
        messages.success(self.request, (success_message))
        return reverse('base:CursoList')


class CursoPdfList(View):
    def get(self, request, *args, **kwargs):
        now = timezone.now()
        data = {
            'section_title': 'Listado de Cursos',
            'module_title': 'Listado de Cursos',
            'today': now,
            'empresa_nombre': EMPRESA_NOMBRE,
            'empresa_rif': EMPRESA_RIF,
            'empresa_direstado': EMPRESA_DIRESTADO,
            'empresa_dirciudad': EMPRESA_DIRCIUDAD,
            'empresa_direccion': EMPRESA_DIRECCION,
            'empresa_postal': EMPRESA_POSTAL,
            'context_pdf': Curso.objects.all()
        }
        pdf = render_to_pdf('base/CursoPdfList.html', data)
        return HttpResponse(pdf, content_type='application/pdf')


class CursoPdfUnid(View):
    def get(self, request, *args, **kwargs):
        id = str
        id = kwargs['pk']
        object_list = Curso.objects.filter(id__icontains=id)
        now = timezone.now()
        data = {
            'section_title': 'Datos del Curso',
            'module_title': 'Datos del Curso',
            'today': now,
            'empresa_nombre': EMPRESA_NOMBRE,
            'empresa_rif': EMPRESA_RIF,
            'empresa_direstado': EMPRESA_DIRESTADO,
            'empresa_dirciudad': EMPRESA_DIRCIUDAD,
            'empresa_direccion': EMPRESA_DIRECCION,
            'empresa_postal': EMPRESA_POSTAL,
            'context_pdf': object_list
        }
        pdf = render_to_pdf('base/CursoPdfUnid.html', data)
        return HttpResponse(pdf, content_type='application/pdf')


class CursoWord(View):
    def get(self, request, *args, **kwargs):
        id = str
        id = kwargs['pk']
        results = Curso.objects.filter(id__exact=id)
        date = datetime.now().strftime('%d-%m-%Y')
        document = Document(os.path.join(BASE_DIR, 'base/base.docx'))
        style_normal = document.styles['Normal']
        font = style_normal.font
        font.name = 'Arial' 
        font.size = Pt(11)
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(1.0)
            section.bottom_margin = Cm(1.0)
            section.left_margin = Cm(1.0)
            section.right_margin = Cm(1.0)
        table_head = document.add_table(rows=3, cols=2)
        row_1 = table_head.rows[0]
        row_1.cells[0].text = EMPRESA_NOMBRE
        formato_emp = row_1.cells[0].paragraphs
        for paragraph in formato_emp:
            for run in paragraph.runs:
                font = run.font
                font.name = 'Arial'
                font.size = Pt(16)
        row_1.cells[1].text = '\t\t\t\tFecha : ' + date
        row_2 = table_head.rows[1]
        row_2.cells[0].text = EMPRESA_RIF
        row_3 = table_head.rows[2]
        row_3.cells[0].text = EMPRESA_DIRCIUDAD + ' ' + EMPRESA_DIRESTADO + ' ' + EMPRESA_POSTAL
        document.add_paragraph()
        title = document.add_paragraph('Datos del Curso')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph('ID : ' + str(results[0].id))
        document.add_paragraph('Descripción : ' +  results[0].descripcion)
        document.add_page_break()
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename={date}-Curso.docx'.format(
            date=datetime.now().strftime('%d-%m-%Y'),
        )
        document.save(response)
        return response


class CursoExcel(View):
    def get(self, request, *args, **kwargs):
        id = str
        id = kwargs['pk']
        results = Curso.objects.filter(id__exact=id)
        date = datetime.now().strftime('%d-%m-%Y')
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
        response['Content-Disposition'] = 'attachment; filename={date}-Curso.xlsx'.format(
            date=datetime.now().strftime('%d-%m-%Y'),
        )
        wb = Workbook()
        ws = wb.active
        ws.column_dimensions['A'].width = 2.0
        ws.column_dimensions['B'].width = 25.0
        ws.column_dimensions['C'].width = 25.0
        ws.column_dimensions['D'].width = 25.0
        ws['B1'] = EMPRESA_NOMBRE
        ws['B2'] = EMPRESA_RIF
        ws['B3'] = EMPRESA_DIRCIUDAD + ' ' + EMPRESA_DIRESTADO
        ws['B4'] = EMPRESA_POSTAL
        ws['D1'] = date
        ws['C5'] = 'Datos del Curso'
        ws['B7'] = 'ID :'
        ws['B8'] = 'Descripción :'
        ws['C7'] = str(results[0].id)
        ws['C8'] = results[0].descripcion
        wb.save(response)
        return response


# Sector:
class SectorList(generic.ListView):
    model = Sector
    template_name = 'base/Sector_i.html'
    queryset = Sector.objects.all().order_by('-created')
    context_object_name = 'Sector'
    extra_context = {
        'section_title': 'Listado de Sectores',
        'module_title': 'Sector'
    }


class SectorCreate(SuccessMessageMixin, CreateView):
    model = Sector
    form_class = SectorForm
    extra_context = {
        'section_title': 'Crear Sector',
        'module_title': 'Sector',
    }
    template_name = 'base/Sector_c.html'
    success_message = 'Creado Correctamente!'

    def get_success_url(self):
        return reverse('base:SectorList')


class SectorDetail(generic.DetailView):
    model = Sector
    form_class = SectorForm
    template_name = 'base/Sector_d.html'

    def get_context_data(self, **kwargs):
        context = super(SectorDetail, self).get_context_data(**kwargs)
        context['now'] = timezone.now()
        extra_context = {
            'section_title': 'Detalle del Sector',
            'module_title': 'Sector'
        }
        return context


class SectorUpdate(SuccessMessageMixin, UpdateView):
    model = Sector
    form_class = SectorForm
    template_name = 'base/Sector_u.html'
    success_message = 'Actualizado Correctamente!'

    def get_success_url(self):
        return reverse('base:SectorList')


class SectorDelete(SuccessMessageMixin, DeleteView):
    model = Sector
    form = Sector
    fields = '__all__'

    def get_success_url(self):
        success_message = 'Eliminado Correctamente!'
        messages.success(self.request, (success_message))
        return reverse('base:SectorList')


class SectorPdfList(View):
    def get(self, request, *args, **kwargs):
        now = timezone.now()
        data = {
            'section_title': 'Listado de Sectores',
            'module_title': 'Listado de Sectores',
            'today': now,
            'empresa_nombre': EMPRESA_NOMBRE,
            'empresa_rif': EMPRESA_RIF,
            'empresa_direstado': EMPRESA_DIRESTADO,
            'empresa_dirciudad': EMPRESA_DIRCIUDAD,
            'empresa_direccion': EMPRESA_DIRECCION,
            'empresa_postal': EMPRESA_POSTAL,
            'context_pdf': Sector.objects.all()
        }
        pdf = render_to_pdf('base/SectorPdfList.html', data)
        return HttpResponse(pdf, content_type='application/pdf')


class SectorPdfUnid(View):
    def get(self, request, *args, **kwargs):
        id = str
        id = kwargs['pk']
        object_list = Sector.objects.filter(id__icontains=id)
        now = timezone.now()
        data = {
            'section_title': 'Datos del Sector',
            'module_title': 'Datos del Sector',
            'today': now,
            'empresa_nombre': EMPRESA_NOMBRE,
            'empresa_rif': EMPRESA_RIF,
            'empresa_direstado': EMPRESA_DIRESTADO,
            'empresa_dirciudad': EMPRESA_DIRCIUDAD,
            'empresa_direccion': EMPRESA_DIRECCION,
            'empresa_postal': EMPRESA_POSTAL,
            'context_pdf': object_list
        }
        pdf = render_to_pdf('base/SectorPdfUnid.html', data)
        return HttpResponse(pdf, content_type='application/pdf')


class SectorWord(View):
    def get(self, request, *args, **kwargs):
        id = str
        id = kwargs['pk']
        results = Sector.objects.filter(id__exact=id)
        date = datetime.now().strftime('%d-%m-%Y')
        document = Document(os.path.join(BASE_DIR, 'base/base.docx'))
        style_normal = document.styles['Normal']
        font = style_normal.font
        font.name = 'Arial'
        font.size = Pt(11)
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(1.0)
            section.bottom_margin = Cm(1.0)
            section.left_margin = Cm(1.0)
            section.right_margin = Cm(1.0)
        table_head = document.add_table(rows=3, cols=2)
        row_1 = table_head.rows[0]
        row_1.cells[0].text = EMPRESA_NOMBRE
        formato_emp = row_1.cells[0].paragraphs
        for paragraph in formato_emp:
            for run in paragraph.runs:
                font = run.font
                font.name = 'Arial'
                font.size = Pt(16)
        row_1.cells[1].text = '\t\t\t\tFecha : ' + date
        row_2 = table_head.rows[1]
        row_2.cells[0].text = EMPRESA_RIF
        row_3 = table_head.rows[2]
        row_3.cells[0].text = EMPRESA_DIRCIUDAD + ' ' + EMPRESA_DIRESTADO + ' ' + EMPRESA_POSTAL
        document.add_paragraph()
        title = document.add_paragraph('Datos del Sector')
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph('ID : ' + str(results[0].id))
        document.add_paragraph('Descripción : ' +  results[0].descripcion)
        document.add_page_break()
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename={date}-Sector.docx'.format(
            date=datetime.now().strftime('%d-%m-%Y'),
        )
        document.save(response)
        return response


class SectorExcel(View):
    def get(self, request, *args, **kwargs):
        id = str
        id = kwargs['pk']
        results = Sector.objects.filter(id__exact=id)
        date = datetime.now().strftime('%d-%m-%Y')
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        )
        response['Content-Disposition'] = 'attachment; filename={date}-Sector.xlsx'.format(
            date=datetime.now().strftime('%d-%m-%Y'),
        )
        wb = Workbook()
        ws = wb.active
        ws.column_dimensions['A'].width = 2.0
        ws.column_dimensions['B'].width = 25.0
        ws.column_dimensions['C'].width = 25.0
        ws.column_dimensions['D'].width = 25.0
        ws['B1'] = EMPRESA_NOMBRE
        ws['B2'] = EMPRESA_RIF
        ws['B3'] = EMPRESA_DIRCIUDAD + ' ' + EMPRESA_DIRESTADO
        ws['B4'] = EMPRESA_POSTAL
        ws['D1'] = date
        ws['C5'] = 'Datos del Sector'
        ws['B7'] = 'ID :'
        ws['B8'] = 'Descripción :'
        ws['C7'] = str(results[0].id)
        ws['C8'] = results[0].descripcion
        wb.save(response)
        return response


#~ Historico de Entrega
class HistoricoEntregaList(SuccessMessageMixin, generic.ListView):
    model = HistoricoEntrega
    template_name = 'base/HistoricoEntrega_i.html'
    queryset = HistoricoEntrega.objects.all().order_by('-created')
    context_object_name = 'HistoricoEntrega'
    extra_context = {
        'section_title': 'Listado de Cajas entregadas a los Beneficiados',
        'module_title': 'Historico Entrega'
    }
    def get(self, request, *args, **kwargs):
        query = self.request.GET.get('q')
        try:
            if query=='ok':
                AsignarCajasMes()
                success_message = 'Listado de Cajas asignado!'

            messages.success(self.request, (success_message))
            return super(HistoricoEntregaList, self).get(request, *args, **kwargs)
        except Exception as ex:
            success_message = 'Error al asignar las cajas del mes!' + str(ex)
            messages.success(self.request, (success_message))
            return super(HistoricoEntregaList, self).get(request, *args, **kwargs)


def AsignarCajasMes():
    listado_beneficiado = Beneficiado.objects.filter(is_active__exact='True')
    i = 0
    for row in listado_beneficiado:
        beneficiado = listado_beneficiado[i]
        beneficiado_id = listado_beneficiado[i].id_beneficiado
        tipo_caja = CajaBeneficiado.objects.filter(beneficiado__exact=beneficiado_id)
        j = 0
        for rows in tipo_caja:
            caja = tipo_caja[j].caja
            new_item = HistoricoEntrega()
            new_item = HistoricoEntrega.objects.create(beneficiado=str(beneficiado), caja=str(caja))
            new_item.save()
            j = j + 1
        i = i + 1
